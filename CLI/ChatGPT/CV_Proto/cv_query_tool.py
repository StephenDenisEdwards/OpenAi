import os
import openai
import docx
import json
from pathlib import Path
from typing import List, Dict
from tqdm import tqdm
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Set your OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")  # or hardcode: openai.api_key = "sk-..."

# Model to use for embeddings
EMBEDDING_MODEL = "text-embedding-3-small"
MAX_INPUT_CHARS = 8000  # Safe max for embedding input

# Function to get embedding from OpenAI API
def get_embedding(text: str, model: str = "text-embedding-3-small") -> List[float]:
    text = text.replace("\n", " ")
    if len(text) > MAX_INPUT_CHARS:
        text = text[:MAX_INPUT_CHARS]
    response = openai.Embedding.create(input=text, model=model)
    return response["data"][0]["embedding"]

# Load and parse CVs from DOCX files
def extract_text_from_docx(file_path: Path) -> str:
    doc = docx.Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    tables = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                tables.append(" | ".join(row_text))

    all_text = paragraphs + tables
    return "\n".join(all_text)

# Load all CVs from a directory
def load_cvs_from_directory(directory: Path) -> List[Dict]:
    cvs = []
    for file in directory.glob("*.docx"):
        text = extract_text_from_docx(file)
        cvs.append({
            "filename": file.name,
            "text": text,
            "embedding": get_embedding(text, model=EMBEDDING_MODEL)
        })
    return cvs

# Save and load embeddings cache (optional)
def save_cv_data(cvs: List[Dict], file: Path):
    with open(file, "w", encoding="utf-8") as f:
        json.dump(cvs, f, ensure_ascii=False, indent=2)

def load_cv_data(file: Path) -> List[Dict]:
    with open(file, "r", encoding="utf-8") as f:
        return json.load(f)

# Ask a question about the CVs
def ask_question(cvs: List[Dict], query: str):
    query_embedding = get_embedding(query, model=EMBEDDING_MODEL)
    ranked = sorted(
        cvs,
        key=lambda x: cosine_similarity([x["embedding"]], [query_embedding])[0][0],
        reverse=True
    )
    top_cv = ranked[0]

    # Feed into GPT for summarised answer
    context = f"""You are a recruiter. Here is a candidate's CV:
{top_cv['text']}
Based on the CV, answer the following question:
{query}
"""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": context}]
    )
    print(f"\nBest match: {top_cv['filename']}")
    print(response.choices[0].message.content)

# Compare two CVs
def compare_cvs(cv1: Dict, cv2: Dict):
    prompt = f"""Compare the following two CVs and highlight the differences in skills, experience, and education.

CV 1 ({cv1['filename']}):
{cv1['text']}

CV 2 ({cv2['filename']}):
{cv2['text']}
"""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    print(response.choices[0].message.content)

# New helper function to ask a question about a specific CV
def ask_question_specific(cv: Dict, query: str):
    # Feed into GPT for summarised answer
    context = f"""You are a recruiter. Here is a candidate's CV:
{cv['text']}
Based on the CV, answer the following question:
{query}
"""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": context}]
    )
    print(f"\nSelected CV: {cv['filename']}")
    print(response.choices[0].message.content)

# Entry point
if __name__ == "__main__":
    base_dir = Path(__file__).parent.resolve()
    cv_dir = base_dir / "cvs"
    cache_file = base_dir / "cv_data.json"

    if cache_file.exists():
        cvs = load_cv_data(cache_file)
    else:
        print("\nLoading and embedding CVs from ./cvs directory...")
        cvs = load_cvs_from_directory(cv_dir)
        save_cv_data(cvs, cache_file)

    while True:
        cmd = input("\nEnter a command (ask / ask_specific / compare / list / exit): ").strip().lower()
        if cmd == "ask":
            question = input("Enter your question: ")
            ask_question(cvs, question)
        elif cmd == "ask_specific":
            print("Available CVs:")
            for idx, cv in enumerate(cvs):
                print(f"  [{idx}] {cv['filename']}")
            idx = int(input("Enter the CV index: "))
            question = input("Enter your question: ")
            ask_question_specific(cvs[idx], question)
        elif cmd == "compare":
            print("Available CVs:")
            for idx, cv in enumerate(cvs):
                print(f"  [{idx}] {cv['filename']}")
            idx1 = int(input("Enter first CV index: "))
            idx2 = int(input("Enter second CV index: "))
            compare_cvs(cvs[idx1], cvs[idx2])
        elif cmd == "list":
            for cv in cvs:
                print(f"- {cv['filename']}")
        elif cmd == "exit":
            break
        else:
            print("Unknown command. Try again.")
